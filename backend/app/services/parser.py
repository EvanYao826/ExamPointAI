"""
Word 文档解析引擎
支持 .docx 格式，提取题目、选项、答案、解析
"""

import re
from docx import Document


def parse_docx(filepath: str) -> list[dict]:
    """
    解析 .docx 文件，返回题目列表

    每个题目格式：
    {
        "content": "题目内容",
        "type": "single_choice",  # single_choice / multiple_choice / judge / fill_blank / short_answer
        "options": [
            {"key": "A", "content": "选项内容", "is_answer": False},
            {"key": "B", "content": "选项内容", "is_answer": True},
            ...
        ],
        "answer": "B",
        "analysis": "解析内容"
    }
    """
    doc = Document(filepath)
    lines = []

    # 提取所有段落文本
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # 合并表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)

    if not lines:
        return []

    # 解析题目
    questions = _parse_questions(lines)
    return questions


def _parse_questions(lines: list[str]) -> list[dict]:
    """解析题目列表"""
    questions = []
    current_question = None
    current_section = None  # track if we're in options, answer, analysis section

    # 题号正则
    question_pattern = re.compile(r'^(\d+)\s*[.、．）\)]\s*(.+)')
    # 选项正则
    option_pattern = re.compile(r'^([A-Fa-f])\s*[.、．）\)]\s*(.+)')
    # 答案正则（单题答案）
    answer_pattern = re.compile(r'^[答案|参考答案]*[：:]\s*([A-Fa-f]+)')
    # 批量答案正则（如 "1.B 2.C 3.A"）
    batch_answer_pattern = re.compile(r'(\d+)\s*[.、．）\)]\s*([A-Fa-f]+)')
    # 解析正则
    analysis_pattern = re.compile(r'^[解析|解题思路]*[：:]\s*(.+)')

    for line in lines:
        # 检查是否是新题目
        q_match = question_pattern.match(line)
        if q_match:
            # 保存上一题
            if current_question:
                _finalize_question(current_question)
                questions.append(current_question)

            # 开始新题目
            q_num = int(q_match.group(1))
            q_content = q_match.group(2).strip()
            current_question = {
                "number": q_num,
                "content": q_content,
                "type": "single_choice",
                "options": [],
                "answer": "",
                "analysis": "",
            }
            current_section = "content"
            continue

        if not current_question:
            continue

        # 检查是否是选项
        opt_match = option_pattern.match(line)
        if opt_match:
            key = opt_match.group(1).upper()
            content = opt_match.group(2).strip()
            current_question["options"].append({
                "key": key,
                "content": content,
                "is_answer": False,
            })
            current_section = "options"
            continue

        # 检查是否是答案行
        ans_match = answer_pattern.match(line)
        if ans_match:
            answer = ans_match.group(1).upper()
            current_question["answer"] = answer
            _mark_answer(current_question["options"], answer)
            current_section = "answer"
            continue

        # 检查是否是批量答案行（如 "参考答案：1.B 2.C 3.A"）
        if "答案" in line and re.search(r'\d+\s*[.、．）\)]\s*[A-Fa-f]', line):
            batch_matches = batch_answer_pattern.findall(line)
            if batch_matches:
                for num_str, ans in batch_matches:
                    num = int(num_str)
                    # 找到对应的题目
                    for q in questions:
                        if q.get("number") == num:
                            q["answer"] = ans.upper()
                            _mark_answer(q["options"], ans.upper())
                continue

        # 检查是否是解析行
        ana_match = analysis_pattern.match(line)
        if ana_match:
            current_question["analysis"] = ana_match.group(1).strip()
            current_section = "analysis"
            continue

        # 续行处理
        if current_section == "content":
            current_question["content"] += "\n" + line
        elif current_section == "analysis":
            current_question["analysis"] += "\n" + line

    # 保存最后一题
    if current_question:
        _finalize_question(current_question)
        questions.append(current_question)

    # 清理临时字段
    for q in questions:
        q.pop("number", None)

    return questions


def _finalize_question(question: dict):
    """完善题目信息"""
    # 判断题型
    if question["options"]:
        if len(question["answer"]) > 1:
            question["type"] = "multiple_choice"
        else:
            question["type"] = "single_choice"
    elif "____" in question["content"] or "___" in question["content"] or "（  ）" in question["content"]:
        question["type"] = "fill_blank"
    elif question["answer"] and question["answer"] in ("对", "错", "√", "×", "T", "F"):
        question["type"] = "judge"
    else:
        question["type"] = "short_answer"

    # 清理内容
    question["content"] = question["content"].strip()
    question["analysis"] = question["analysis"].strip()


def _mark_answer(options: list[dict], answer: str):
    """标记正确答案"""
    for opt in options:
        if opt["key"] in answer:
            opt["is_answer"] = 1
